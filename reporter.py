from docx import *
from docx.shared import *
from pathlib import Path
from os import system, name, path, makedirs, getuid, listdir
from os.path import isfile, join
import subprocess


class Reporter():
    def __init__(self, name, host, ports, services, tabs):
        self.name = name
        self.host = host
        self.ports = ports
        self.services = services
        self.tabs = tabs

    def check_line(self, s1, s2, lines):
     for line in lines:
         if s1 in line and s2 in line:
             return True
     return False


    def find_in_par(self, txt,doc):               # FIND paragraph containing some txt
        for paragraph in doc.paragraphs:
            if txt in paragraph.text:
                return paragraph
        return None

    def find_second_in_par(self, txt,doc):               # FIND paragraph containing some txt
        x=False
        for paragraph in doc.paragraphs:
            if txt in paragraph.text and not x:
                x=True
            elif txt in paragraph.text and x:
                return paragraph
        return None

    def clear_table(self, table):
        i=0
        for p in table.cell(0,0).paragraphs:
            if i==0:
                i=1
            else:
                self.delete_paragraph(p)
        return table

    def find_in_tab(self, txt,doc):               # FIND paragraph containing some txt
        for table in doc.tables:
            if txt in table.cell(0,0).paragraphs[1].text:
                return table
        return None


    def index_before_par(self, txt,doc):               # FIND paragraph containing some txt
        for i in range(0,len(doc.paragraphs)-1):
            if txt in doc.paragraphs[i].text:
                return i-1
        return -1
        
    def valid_xml_char_ordinal(self, c):
        codepoint = ord(c)
        # conditions ordered by presumed frequency
        return (
            0x20 <= codepoint <= 0xD7FF or
            codepoint in (0x9, 0xA, 0xD) or
            0xE000 <= codepoint <= 0xFFFD or
            0x10000 <= codepoint <= 0x10FFFF
            )

    def sanitize(self, s):
        return ''.join(c for c in s if self.valid_xml_char_ordinal(c))

    def delete_paragraph(self, paragraph):        # DELETE chosen paragraph
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    # table._element.getparent().remove(table._element)     # DELETE table


    def merge_docs(self, files):          # MERGE documents # files is a list of documents to merge
        merged_document = Document()


        for index, f in enumerate(files):
            sub_doc = Document(f)

            # Don't add a page break if you've reached the last file.
            if index < len(files)-1:
               sub_doc.add_page_break()

            for element in sub_doc.element.body:
                merged_document.element.body.append(element)

        return merged_document

    def merge_before_par_docs(self, d, sub_doc):          # MERGE documents before SUGGESTED SOLUTIONS  # tup is a tuple of two documents to merge

        ind=self.index_before_par("SUGGESTED SOLUTIONS",d)

        par=d.paragraphs[ind]

        for element in sub_doc.element.body:
            par.element.body.append(element)

        return d

    def IIS_isvuln(self, IISouts):
        if IISouts == []:
            return None
        for o in IISouts:
            if "Result: Not vulnerable" not in o[1]:
                return o
        return None

    def LRH_isvuln(self, LRHouts):
        if LRHouts == []:
            return None
        for o in LRHouts:
            if "<*>" in o[1]:
                return o
        return None
    
    def grep(self, s,text):
        lista=text.split("\n")
        for l in lista:
            if s in l:
                return l
        return ""
    
    def gen_SECMISC(self, olrh, oiis):      # check if SEC-MISC and its content is True
        doc = Document('Templates/SEC-MISC.docx')

        if oiis==None:
            # deleting 8.3 enum section
            par=self.find_in_par("file names supported by the HTTP service and accepted by the system:",doc)
            self.delete_paragraph(par)

            par=self.find_in_par("8dot3 ENUMERATION",doc)
            self.delete_paragraph(par)

            par=self.find_in_par("During the analysis it was found that the NTFS file system in use on the operating system ",doc)
            self.delete_paragraph(par)

            par=self.find_in_par("– 8.3 enumeration result",doc)
            self.delete_paragraph(par)

            table=self.find_in_tab("java -jar iis_shortname_scanner.jar", doc)
            table._element.getparent().remove(table._element)
        else:
            table=self.find_in_tab("java -jar iis_shortname_scanner.jar", doc)
            table=self.clear_table(table)
            table.cell(0,0).paragraphs[0].text=self.sanitize(oiis[1]).strip()

                

        if olrh==None:
            # deleting HTTP methods section
            par=self.find_in_par("INSECURE HTTP VERBS",doc)
            self.delete_paragraph(par)

            par=self.find_in_par("HTTP methods considered insecure or unnecessary that have not been disabled ",doc)
            self.delete_paragraph(par)

            par=self.find_in_par("- List of used and recognized methods",doc)
            self.delete_paragraph(par)

            table=self.find_in_tab("Allow: GET, POST, OPTIONS, HEAD, MKCOL", doc)
            table._element.getparent().remove(table._element)
        else:
            table=self.find_in_tab("Allow: GET, POST, OPTIONS, HEAD, MKCOL", doc)
            table=self.clear_table(table)
            table.cell(0,0).paragraphs[0].text=self.sanitize(self.grep("<*>",olrh[1])).strip()

        
        doc.save('Reports/Temp/2-SEC-MISC.docx')
        return
    
    def gen_SSAPNU(self, RHouts, LRHouts, host):      # check if SSAP-NU and its content is True
        doc = Document('Templates/SSAP-NU.docx')
        
        if LRHouts == []:
            par=self.find_in_par("The next Box shows the evidence of the software component version in the response header.",doc)
            self.delete_paragraph(par)

            par=self.find_in_par(" – Leak software version: not updated",doc)
            self.delete_paragraph(par)

            table=self.find_in_tab("Response: https://10.72.16.13:2224/", doc)
            table._element.getparent().remove(table._element)
            
        #else:
        #    for o in LRHouts:
        #        if 
        i=0
        for o in RHouts:
            if i==0:
                # modify current doc
                ou=o[1].strip()
                ou = ''.join(c for c in ou if self.valid_xml_char_ordinal(c))
                i=1
                first=ou.split("\n")[0]
                second=ou.split("\n")[1]

                par=self.find_in_par("SoftwareName",doc)
                par.text=par.text.replace("SoftwareName",first.split("/")[0])

                #par=self.find_in_par("2.4.29",doc)
                par.text=par.text.replace("104.214.239.16",host)
                par.text=par.text.replace("2.4.29",first.split("/")[1].split(":")[0])

                #par=self.find_second_in_par("HIGH",doc)
                par.text=par.text.replace("HIGH",first[first.find("(")+1:first.find(")")].upper())

                #par=self.find_in_par("(CVE-2019-0211)",doc)
                par.text=par.text.replace("(CVE-2019-0211)","("+first[first.find(":")+1:first.find("(")].upper()+")")

                #par=self.find_in_par("CVE-2017-15710…",doc)
                par.text=par.text.replace("CVE-2017-15710…",second)

                if LRHouts == []:
                    par=self.find_in_par("The next Box shows the evidence of the software component version in the response header.",doc)
                    self.delete_paragraph(par)

                    par=self.find_in_par(" – Leak software version: not updated",doc)
                    self.delete_paragraph(par)

                    table=self.find_in_tab("Response: https://10.72.16.13:2224/", doc)
                    table._element.getparent().remove(table._element)
                else:
                    for lrh in LRHouts:
                        if first.split("/")[0] in lrh[1].lower() and first.split("/")[1].split(":")[0].strip() in lrh[1].lower():
                            table=self.find_in_tab("Response: https://10.72.16.13:2224/", doc)
                            table=self.clear_table(table)
                            table.cell(0,0).paragraphs[0].text=self.sanitize(self.extract_head(lrh[1])).strip()                                       
                            break
                            
                doc.save('Reports/Temp/3-SSAP-NU.docx')

            else:
                # modify elem and append
                doc2 = Document('Templates/SSAP-NU-elem.docx')
                ou=o[1].strip()
                ou = ''.join(c for c in ou if self.valid_xml_char_ordinal(c))
                i=1
                first=ou.split("\n")[0]
                second=ou.split("\n")[1]

                par=self.find_in_par("SoftwareName",doc2)
                par.text=par.text.replace("SoftwareName",first.split("/")[0])

                #par=self.find_in_par("2.4.29",doc)
                par.text=par.text.replace("104.214.239.16",host)
                par.text=par.text.replace("2.4.29",first.split("/")[1].split(":")[0])

                #par=self.find_second_in_par("HIGH",doc)
                par.text=par.text.replace("HIGH",first[first.find("(")+1:first.find(")")].upper())

                #par=self.find_in_par("(CVE-2019-0211)",doc)
                par.text=par.text.replace("(CVE-2019-0211)","("+first[first.find(":")+1:first.find("(")].upper()+")")

                #par=self.find_in_par("CVE-2017-15710…",doc)
                par.text=par.text.replace("CVE-2017-15710…",second)
                
                if LRHouts == []:
                    par=self.find_in_par("The next Box shows the evidence of the software component version in the response header.",doc2)
                    self.delete_paragraph(par)
        
                    par=self.find_in_par(" – Leak software version: not updated",doc2)
                    self.delete_paragraph(par)
        
                    table=self.find_in_tab("Response: https://10.72.16.13:2224/", doc2)
                    table._element.getparent().remove(table._element)
                else:
                    for lrh in LRHouts:
                        if first.split("/")[0] in lrh[1].lower() and first.split("/")[1].split(":")[0].strip() in lrh[1].lower():
                            table=self.find_in_tab("Response: https://10.72.16.13:2224/", doc2)
                            table=self.clear_table(table)
                            table.cell(0,0).paragraphs[0].text=self.sanitize(self.extract_head(lrh[1])).strip()
                            break

                doc2.save('Reports/Temp/doc2.docx')
                doc = self.merge_docs(['Reports/Temp/3-SSAP-NU.docx','Reports/Temp/doc2.docx'])
                doc.save('Reports/Temp/3-SSAP-NU.docx')
                bashCommand = "rm Reports/Temp/doc2.docx"
                process = subprocess.Popen(bashCommand.split(), stdout=subprocess.PIPE)
                process.communicate()
        doc.add_page_break()
        doc.save('Reports/Temp/3-SSAP-NU.docx')
        return
    
    def gen_NOCRYPTT(self, ports, services):      # check if NOCRYPTT is True
        doc = Document('Templates/NOCRYPTT.docx')
        doc.save('Reports/Temp/4-NOCRYPTT.docx')
        return
    
    def gen_CRYPWK(self, TSSLouts):      # check if CRYP-WK is True
        doc = Document('Templates/CRYP-WK.docx')
        table=self.find_in_tab("Output testssl", doc)
        table=self.clear_table(table)
        table.cell(0,0).paragraphs[0].text=self.sanitize(TSSLouts[0][1]).strip()
        doc.save('Reports/Temp/5-CRYP-WK.docx')
        return
    
    def gen_RESPH(self, LRHouts):      # check if RESP-H is True
        for lrh in LRHouts:
            if "<!>" in lrh[1].lower():
                doc = Document('Templates/RESP-H.docx')
                table=self.find_in_tab('RESPONSE https://11.22.33.44/', doc)
                table=self.clear_table(table)
                table.cell(0,0).paragraphs[0].text=self.sanitize(self.extract_head(lrh[1])).strip()
                doc.save('Reports/Temp/6-RESP-H.docx')
                break
        return
    
    def gen_CJKPROT(self, LRHouts):      # check if CJK-PROT is True
        for lrh in LRHouts:
            if self.check_line("<!>", "X-Frame-Options", lrh[1].lower().split("\n")):       # se sono nella stessa riga
                doc = Document('Templates/CJK-PROT.docx')
                table=self.find_in_tab('RESPONSE', doc)
                table=self.clear_table(table)
                table.cell(0,0).paragraphs[0].text=self.sanitize(self.extract_head(lrh[1])).strip()                                 
                doc.save('Reports/Temp/7-CJK-PROT.docx')
                break
        return
    
    def gen_SCKPROT(self, LRHouts):      # check if SCK-PROT is True
        for lrh in LRHouts:
            if "<*>" in lrh[1].lower():     
                doc = Document('Templates/SCK-PROT.docx')
                table=self.find_in_tab('RESPONSE', doc)
                table=self.clear_table(table)
                table.cell(0,0).paragraphs[0].text=self.sanitize(self.extract_head(lrh[1])).strip()
                doc.save('Reports/Temp/8-SCK-PROT.docx')
                break      
        return
    
    def gen_INLEAK(self, LRHouts):      # check if IN-LEAK is True
        for lrh in LRHouts:
            if "<?>" in lrh[1].lower():
                doc = Document('Templates/IN-LEAK.docx')
                table=self.find_in_tab('RESPONSE https://11.22.33.44/', doc)
                table=self.clear_table(table)
                table.cell(0,0).paragraphs[0].text=self.sanitize(self.extract_head(lrh[1])).strip()
                doc.save('Reports/Temp/9-IN-LEAK.docx')
                break        
        return

    def extract_head(self,out):
        ret=''
        y=False
        for line in out.split("\n"):
            if "###############################" in line:
                break
            if "~~~" in line:
                ret+=self.sanitize(line).strip().strip("~~~").strip()
                y=True
            elif y==True:
                ret+=self.sanitize(line).strip()+"\n"
        return ret
    
    def gen_TECHANN(self, ports, services):      # check if TECH-ANN is True
        doc = Document('Templates/TECH-ANN.docx')
        doc.save('Reports/Temp/TECH-ANN.docx')
        return
    

    def caller_gen(self, host, ports, services, tabs):
        LRHouts=[]
        RHouts=[]
        TSSLouts=[]
        IISouts=[]
        for X in tabs:
            if X['tool']=='IIS Shortname Scanner':
                IISouts=X['output']
            elif X['tool']=='LiteRespH':
                LRHouts=X['output']
            elif X['tool']=='RHsecapi':
                RHouts=X['output']
            elif X['tool']=='TestSSL.sh':
                TSSLouts=X['output']


            if LRHouts != []:
                self.gen_RESPH(LRHouts)
                self.gen_CJKPROT(LRHouts)
                self.gen_SCKPROT(LRHouts)
                self.gen_INLEAK(LRHouts)
            
            if RHouts != []:
                self.gen_SSAPNU(RHouts,LRHouts, host)

            if TSSLouts != []:
                self.gen_CRYPWK(TSSLouts)

            if "http" in services:
                self.gen_NOCRYPTT(ports, services)

            oiis=self.IIS_isvuln(IISouts)
            olrh=self.LRH_isvuln(LRHouts)
            
            if oiis!=None or olrh!=None:
                self.gen_SECMISC(olrh, oiis)

        return

    

    def generate_report(self):
        Path("Reports/Temp").mkdir(parents=True, exist_ok=True)

        self.caller_gen(self.host, self.ports, self.services, self.tabs)

        files=[]

        mypath="Reports/Temp/"

        # for f in get_files().sort():
        #   files.append(f)

        bashCommand = "ls " + mypath
        process = subprocess.Popen(bashCommand.split(), stdout=subprocess.PIPE)
        o,e = process.communicate()
        files = [join(mypath, f) for f in o.decode().split("\n") if ".docx" in f]

        merged_document = self.merge_docs(files)
        for i in range(0,len(merged_document.paragraphs)-1):
            if "System software and " in merged_document.paragraphs[i].text:
                self.delete_paragraph(merged_document.paragraphs[i-1])
                self.delete_paragraph(merged_document.paragraphs[i-2])
                self.delete_paragraph(merged_document.paragraphs[i-3])
                break
        #merged_document.paragraphs[3].text=''
        path='Reports/'+self.name+'.docx'
        merged_document.save(path)
        for f in files:     # clean Temp
            bashCommand = "rm " + f
            process = subprocess.Popen(bashCommand.split(), stdout=subprocess.PIPE)
            process.communicate()
        return path
